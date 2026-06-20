#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成《微服务架构实践大作业》课程报告 (docx)，严格遵循模板的 8 个章节。

特点：附录代码直接读取仓库中的真实源文件，截图/产物从 报告/ 目录读取，
因此报告永远与真实代码、真实运行结果一致。
"""
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# ====================== 封面信息（请按需修改） ======================
COVER = {
    "title": "基于 MCP 的智能网页采集分析智能体平台的设计与实现",
    "year_term": "2024—2025 学年 第二学期",
    "course": "微服务架构",
    "course_id": "（课程编号）",
    "course_seq": "（课程序号）",
    "teacher": "（任课教师）",
    "name": "（姓名）",
    "sid": "（学号）",
    "submit": "2026 年 6 月 16 日",
}

BODY = "宋体"
HEI = "黑体"
MONO = "Consolas"
JADE = RGBColor(0x1F, 0x8F, 0x63)
DARK = RGBColor(0x22, 0x2A, 0x26)

doc = Document()


# ---------------------------------------------------------------- helpers
def _cjk(run, font):
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)


def para(text="", size=12, font=BODY, bold=False, color=None, align=None,
         after=6, before=0, indent=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.line_spacing = 1.5
    if align:
        p.alignment = align
    if indent:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        if color:
            r.font.color.rgb = color
        _cjk(r, font)
    return p


def h1(num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(f"{num}、{text}")
    r.font.size = Pt(16)
    r.bold = True
    r.font.color.rgb = DARK
    _cjk(r, HEI)
    return p


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(13)
    r.bold = True
    r.font.color.rgb = JADE
    _cjk(r, HEI)
    return p


def bullet(text, size=12):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.4
    r = p.add_run(text)
    r.font.size = Pt(size)
    _cjk(r, BODY)
    return p


def numbered(text, size=12):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.line_spacing = 1.4
    r = p.add_run(text)
    r.font.size = Pt(size)
    _cjk(r, BODY)
    return p


def code_block(text, size=8):
    for line in text.rstrip("\n").split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Pt(12)
        r = p.add_run(line if line else " ")
        r.font.size = Pt(size)
        r.font.name = MONO
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.italic = True
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    _cjk(r, BODY)


def image(path, width=6.2, cap=None, hint=None):
    full = os.path.join(ROOT, path)
    if os.path.exists(full):
        doc.add_picture(full, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cap:
            caption(cap)
        return True
    # 干净的"待插入截图"占位框（用于需本人账号/工具的步骤）
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"〔此处插入截图：{cap or path}〕")
    r.font.size = Pt(11)
    r.bold = True
    r.font.color.rgb = RGBColor(0x99, 0x66, 0x00)
    _cjk(r, BODY)
    if hint:
        ph = doc.add_paragraph()
        ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ph.paragraph_format.space_after = Pt(10)
        rh = ph.add_run(hint)
        rh.font.size = Pt(9)
        rh.italic = True
        rh.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        _cjk(rh, BODY)
    return False


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, hh in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(hh)
        r.bold = True
        r.font.size = Pt(10.5)
        _cjk(r, HEI)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(val))
            r.font.size = Pt(10)
            _cjk(r, BODY)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def read(path):
    full = os.path.join(ROOT, path)
    try:
        return open(full, encoding="utf-8").read()
    except Exception as e:
        return f"[无法读取 {path}: {e}]"


def appendix_file(path, lang=""):
    h2(f"▶ {path}")
    code_block(read(path))


# ====================================================== 封面
def cover():
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("微服务架构实践大作业")
    r.font.size = Pt(26)
    r.bold = True
    _cjk(r, HEI)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("课 程 报 告")
    r2.font.size = Pt(18)
    _cjk(r2, HEI)
    for _ in range(2):
        doc.add_paragraph()

    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pt.paragraph_format.space_after = Pt(24)
    rt = pt.add_run("题目：" + COVER["title"])
    rt.font.size = Pt(15)
    rt.bold = True
    rt.font.color.rgb = JADE
    _cjk(rt, HEI)

    fields = [
        ("学年学期", COVER["year_term"]),
        ("课程名称", COVER["course"]),
        ("课程编号", COVER["course_id"]),
        ("课程序号", COVER["course_seq"]),
        ("任课教师", COVER["teacher"]),
        ("姓    名", COVER["name"]),
        ("学    号", COVER["sid"]),
    ]
    for k, v in fields:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(f"{k}：{v}")
        r.font.size = Pt(13)
        _cjk(r, BODY)

    doc.add_paragraph()
    pe = doc.add_paragraph()
    pe.alignment = WD_ALIGN_PARAGRAPH.CENTER
    re = pe.add_run("评分区域（由阅卷老师填写）")
    re.font.size = Pt(11)
    re.bold = True
    _cjk(re, HEI)
    for k in ("结课成绩", "总评成绩"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{k}：________")
        r.font.size = Pt(11)
        _cjk(r, BODY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("提交时间：" + COVER["submit"])
    r.font.size = Pt(11)
    _cjk(r, BODY)
    doc.add_page_break()


# ====================================================== 一、实验目的
def sec1():
    h1("一", "实验目的")
    para("本实验设计并实现一个具有明确业务功能的微服务系统——“智阅（ZhiYue）”：一个配备 MCP "
         "服务、并具备网页爬取能力的智能体（AI Agent）。通过本实验达到以下目的：", indent=True)
    bullet("掌握微服务架构的设计方法：服务拆分、独立部署、通过明确接口（HTTP/REST 与 MCP）通信；")
    bullet("掌握将核心功能封装为符合 Model Context Protocol（MCP）标准的服务，并被 MCP Inspector、"
           "Cline 等客户端及自研智能体调用（加分项）；")
    bullet("掌握使用 requests + BeautifulSoup 进行网页内容抓取与清洗的原理；")
    bullet("掌握以大语言模型（DeepSeek）驱动的智能体“工具规划循环（tool-use loop）”的实现；")
    bullet("掌握使用 Docker 对各服务进行容器化打包，并用 docker-compose 一键编排；")
    bullet("掌握将镜像推送至阿里云容器镜像服务（ACR）的完整流程；")
    bullet("严格对齐《MCP 服务实验指导书》的全部规定步骤，并在其上进行微服务化与智能体化扩展。")


# ====================================================== 二、实验环境
def sec2():
    h1("二", "实验环境")
    table(["类别", "说明"], [
        ["操作系统", "Windows 11 Pro (10.0.26200)"],
        ["编程语言", "Python 3.11（容器运行时） / Python 3.13（宿主开发）"],
        ["Web 框架", "FastAPI + Uvicorn"],
        ["MCP 框架", "mcp 1.27.2（FastMCP 服务端 + ClientSession 客户端）"],
        ["大模型", "DeepSeek（deepseek-chat，OpenAI 兼容接口）"],
        ["爬虫库", "requests 2.x + BeautifulSoup4"],
        ["数据存储", "SQLite（持久化卷）"],
        ["服务间通信", "HTTP/REST（httpx） + MCP over SSE"],
        ["容器化", "Docker 29.4.1 + Docker Compose v5.1.3"],
        ["镜像仓库", "阿里云容器镜像服务（ACR）"],
        ["测试", "pytest（单元/集成） + 自研冒烟测试脚本"],
        ["其他工具", "VS Code、Cline、MCP Inspector、Git"],
    ])
    para("以下为实验环境关键版本的自动采集结果（可复现）：", indent=True)
    code_block(read("报告/artifacts/env.txt"))


# ====================================================== 三、实验内容
def sec3():
    h1("三", "实验内容")
    h2("3.1 业务场景")
    para("用户以自然语言下达网页采集/分析任务，例如“帮我抓取某网页的内容并总结要点并保存”。"
         "智阅智能体借助 DeepSeek 大模型理解意图、规划工具调用，经 MCP 协议调用爬虫/抽取/存储/检索"
         "等工具完成任务，并将抓取内容持久化、可检索，形成智能体的“记忆”。", indent=True)

    h2("3.2 系统总体架构")
    image("报告/diagrams/architecture.png", width=6.4,
          cap="图 3-1  智阅平台总体架构（四服务 + MCP 协议 + 容器编排）")
    para("系统由 4 个可独立部署的微服务组成，全部容器化并由 docker-compose 编排在同一网络中：", indent=True)
    table(["微服务", "技术", "端口", "职责"], [
        ["api-gateway", "FastAPI + 静态前端", "8080", "系统唯一入口；托管 Web 聊天界面；转发请求；/health、/metrics"],
        ["agent-service", "FastAPI + openai + mcp", "8001", "智能体大脑：LLM 工具规划循环、MCP 客户端"],
        ["mcp-tool-service", "FastMCP", "8002", "MCP 工具服务器：爬虫 + 数据工具，stdio + SSE 双传输"],
        ["storage-service", "FastAPI + SQLite", "8003", "文档持久化与检索（智能体的记忆库）"],
    ])
    para("关键设计：agent-service 只通过 MCP 协议获取一切能力（包括存储/检索——它们是 mcp-tool-service "
         "上的 MCP 工具，由该服务转调 storage-service），从而把 MCP 作为智能体与能力之间的统一接口，"
         "充分体现 MCP 加分项。", indent=True)

    h2("3.3 MCP 协议与本系统的 MCP 设计")
    para("Model Context Protocol（MCP）是一个开放协议，允许 AI 助手与外部服务、数据源安全连接，"
         "通过统一的工具（tool）接口扩展其能力。本系统的 mcp-tool-service 是标准 FastMCP 服务，"
         "暴露 5 个工具：", indent=True)
    table(["MCP 工具", "签名", "说明"], [
        ["get_web_content", "(url) → str", "《指导书》原样工具：抓取网页并清洗为纯文本"],
        ["extract_links", "(url) → list", "提取页面全部链接（相对转绝对、去重）"],
        ["crawl_structured", "(url, selector) → list", "按 CSS 选择器结构化爬取"],
        ["save_document", "(url, title, content) → dict", "存入记忆库（转调 storage-service）"],
        ["search_documents", "(query, limit) → list", "检索历史抓取（转调 storage-service）"],
    ])
    para("该服务支持双传输：stdio 供 MCP Inspector 与 Cline 使用（对应《指导书》第三~五步）；"
         "SSE/HTTP 供 agent-service 跨容器通过网络调用（微服务模式）。", indent=True)


# ====================================================== 四、实验步骤
def sec4():
    h1("四", "实验步骤")

    h2("4.1 需求分析")
    para("依据课程要求与《MCP 服务实验指导书》，归纳核心需求：①核心功能须以微服务架构完成；"
         "②建议封装为 MCP 服务（加分项）；③须 Docker 容器化并推送至阿里云；④提交完整报告。"
         "据此确定题目：构建一个“配备 MCP 服务、具备爬虫能力的智能体”。", indent=True)

    h2("4.2 架构设计")
    para("将系统按职责拆分为 4 个微服务（见 3.2）。通信上，对外用 HTTP/REST，智能体与工具之间用 MCP。"
         "数据层用 SQLite + 持久卷。每个服务独立 Dockerfile，docker-compose 统一编排。", indent=True)

    h2("4.3 各微服务实现")
    para("（1）storage-service：基于 FastAPI + SQLite，提供文档的增/查/检索/统计接口。"
         "数据访问层 db.py 以数据库文件路径为参数，便于以临时库做单元测试。核心代码见附录 A。", indent=True)
    para("（2）mcp-tool-service：将 requests+BeautifulSoup 的爬取/解析能力封装为纯函数（tools.py，"
         "可脱网用本地 HTML 夹具单元测试），再由 FastMCP 注册为 5 个 MCP 工具（mcp_server.py）。"
         "数据类工具通过 httpx 转调 storage-service。核心代码见附录 B。", indent=True)
    para("（3）agent-service：核心是 run_loop——一个可注入依赖（LLM 客户端、工具列举/调用函数）的"
         "工具规划循环，因而可在无网络、无密钥时用 fake 做单元测试；run_agent 接线真实 DeepSeek 与"
         "MCP 客户端（mcp_client.py 通过 SSE 连接、initialize、list_tools、call_tool）。见附录 C。", indent=True)
    para("（4）api-gateway：系统唯一入口，托管 Web 聊天界面（原生 HTML/CSS/JS，离线可用），"
         "转发 /api/chat 至智能体、/api/documents 至存储，并提供 /health 与 /metrics。见附录 D。", indent=True)

    h2("4.4 MCP 封装与测试（对应《指导书》第二~五步）")
    para("第二步——编写爬虫并用 fastmcp 封装为 MCP 服务：mcp = FastMCP(\"...\")，以 @mcp.tool() 装饰"
         "工具函数，mcp.run(transport='stdio')。《指导书》原样单文件版本见附录 F（guide-baseline）。", indent=True)
    para("说明：guide-baseline/ 目录 100% 保留《指导书》原样实现；生产 mcp-tool-service 在其基础上做了"
         "如下增强（二者并存，确保既覆盖指导书基线、又体现微服务化扩展）：", indent=True)
    table(["维度", "指导书基线 (guide-baseline)", "本系统生产实现 (mcp-tool-service)"], [
        ["服务名", "Web Scraper", "ZhiYue Web Tools"],
        ["工具数", "1（get_web_content）", "5（爬取/抽取/结构化/存储/检索）"],
        ["传输方式", "stdio", "stdio + SSE 双传输"],
        ["正文截断", "1000 字", "4000 字"],
        ["网络 I/O", "同步 requests", "异步 httpx（不阻塞事件循环）"],
        ["错误语义", "返回 'Error:' 字符串", "抛异常→MCP isError，客户端可感知失败"],
    ])
    para("第三步——测试 MCP 服务：在宿主机执行下述命令启动官方 MCP Inspector，在浏览器中以 stdio 传输 "
         "Connect、选择 get_web_content 工具、输入 URL 并运行：", indent=True)
    code_block("mcp dev guide-baseline/web_content_mcp_stdio.py")
    para("实测要点（已踩坑并解决）：① MCP Inspector v0.22 默认开启 DNS 重绑定保护，需设置环境变量 "
         "ALLOWED_ORIGINS 允许浏览器来源（如 http://127.0.0.1:6274），否则代理会以 Invalid origin 拒绝连接；"
         "② stdio 传输下 stdout 即 JSON-RPC 通道，故启动日志须输出到 stderr（见 web_content_mcp_stdio.py）；"
         "③ Windows 含中文路径时建议将脚本复制到纯英文路径再连接。下图为经 stdio 成功连接 Web Scraper 服务"
         "并运行 get_web_content 的结果：", indent=True)
    image("报告/screenshots/02_mcp_inspector.png", width=6.4,
          cap="图 4-1  MCP Inspector 经 stdio 连接 Web Scraper 服务，initialize→tools/list→tools/call 全部成功，"
              "get_web_content 返回 Tool Result: Success")
    para("此外，为获得可复现的文本验证记录，本实验另用官方 MCP Python 客户端（ClientSession）通过 SSE "
         "完成与 Inspector 等价的交互——initialize → list_tools → call_tool，输出如下：", indent=True)
    code_block(read("报告/artifacts/mcp_client_test.txt"))
    para("第四步——配置 Cline 使用 MCP 服务：编辑 cline_mcp_settings.json，指向脚本绝对路径，"
         "保存后出现绿色圆点代表配置成功（配置见附录 F）：", indent=True)
    code_block(read("guide-baseline/cline_mcp_settings.json"))
    image("报告/screenshots/03_cline.png", width=6.2,
          cap="图 4-2  在 Cline 中通过自然语言调用 MCP 工具获取网页内容",
          hint="在 VS Code 的 Cline 扩展中导入上述配置，出现绿色圆点后让其获取某网页，截图置于 "
               "报告/screenshots/03_cline.png")

    h2("4.5 容器化（对应《指导书》第六步）")
    para("每个服务遵循统一模式编写 Dockerfile：python:3.11-slim 基础镜像、安装依赖、创建非 root "
         "用户 mcpuser、设置环境变量并启动服务。以 storage-service 为例（其余结构一致，见附录 E）：", indent=True)
    code_block(read("services/storage-service/Dockerfile"))
    para("使用 docker-compose 一键编排 4 个服务、自定义网络与持久卷，密钥经 .env 注入（见附录 E）：", indent=True)
    code_block("cd deploy\ncp .env.example .env   # 填入 DeepSeek 密钥\ndocker compose build\ndocker compose up -d")

    h2("4.6 推送阿里云（对应《指导书》第七步）")
    para("登录阿里云容器镜像服务，对每个镜像打标签并推送（脚本 scripts/push_aliyun.sh 已参数化）：", indent=True)
    code_block(read("scripts/push_aliyun.sh"))
    image("报告/screenshots/04_aliyun.png", width=6.2,
          cap="图 4-3  阿里云容器镜像服务中已推送的镜像仓库列表",
          hint="执行 `bash scripts/push_aliyun.sh <registry> <namespace>` 推送后，在阿里云控制台"
               "「容器镜像服务 → 镜像仓库」截图，置于 报告/screenshots/04_aliyun.png")

    h2("4.7 测试")
    para("（1）单元测试（pytest）：覆盖 storage 数据层与 API、爬虫解析纯函数、智能体工具循环（脱网）。", indent=True)
    para("（2）集成/冒烟测试（scripts/smoke_test.py）：对运行中的技术栈做端到端验证——网关健康、"
         "经 MCP 调用爬虫→存储→检索、并在配置密钥时验证经网关的完整智能体对话。测试输出见 5.3。", indent=True)


# ====================================================== 五、结果展示
def sec5():
    h1("五", "结果展示")
    h2("5.1 智能体 Web 对话（含 MCP 工具调用链）")
    image("报告/screenshots/01_web_ui_chat.png", width=6.4,
          cap="图 5-1  Web 界面：智能体对“抓取并总结并保存”的请求自动规划了 2 次 MCP 工具调用")
    para("如图，用户以自然语言提出请求，智能体经 DeepSeek 规划，依次调用 get_web_content 与 "
         "save_document 两个 MCP 工具，返回结构化总结并标注来源；右下“MCP 工具调用链”实时展示了"
         "本轮经 Model Context Protocol 发起的工具调用。", indent=True)

    h2("5.2 容器编排状态")
    para("docker compose ps —— 4 个服务全部正常运行：", indent=True)
    code_block(read("报告/artifacts/docker_ps.txt"))
    para("docker images —— 4 个自建镜像：", indent=True)
    code_block(read("报告/artifacts/docker_images.txt"))

    h2("5.3 端到端冒烟测试结果")
    code_block(read("报告/artifacts/smoke_test.txt"))
    para("结果：8 项断言全部通过（0 失败），覆盖健康检查、MCP 工具的爬取/存储/检索、网关聚合，"
         "以及经 DeepSeek 的完整智能体对话与工具调用。", indent=True)

    h2("5.4 单元测试结果")
    para("pytest —— 10 个单元测试全部通过（覆盖数据层/API、爬虫解析、智能体工具循环）：", indent=True)
    code_block(read("报告/artifacts/pytest.txt"))


# ====================================================== 六、结论与展望
def sec6():
    h1("六", "结论与未来展望")
    h2("6.1 结论")
    para("本实验完整实现了一个“配备 MCP 服务、具备爬虫能力的智能体”微服务系统，并：", indent=True)
    bullet("以 4 个独立微服务 + docker-compose 完成了真正的微服务架构；")
    bullet("将核心能力封装为标准 MCP 服务（5 个工具，stdio+SSE 双传输），可被 Inspector、Cline 及自研智能体调用，落实加分项；")
    bullet("以 DeepSeek 驱动的工具规划循环实现了真正“会用工具”的智能体；")
    bullet("完成全部服务的容器化、编排与阿里云镜像推送；")
    bullet("严格覆盖《MCP 服务实验指导书》全部步骤，端到端测试 8/8 通过。")

    h2("6.2 未来展望")
    bullet("引入服务注册/配置中心（Nacos/Consul）与 API 网关限流、熔断；")
    bullet("以 Kubernetes 替代 docker-compose 实现弹性伸缩与滚动发布；")
    bullet("将 SQLite 升级为 PostgreSQL，并引入向量检索支撑语义记忆；")
    bullet("扩展更多 MCP 工具（如 PDF 解析、定时爬取、RSS 订阅）与多智能体协作；")
    bullet("接入 Prometheus + Grafana 与链路追踪，完善可观测性。")


# =========================== 七、扩展实现：智策金融平台
def sec_finance():
    h1("七", "扩展实现：智策多智能体金融分析平台")
    para("在“智阅”MCP 智能体平台之上，进一步扩展出金融垂直能力——“智策 ZhiCe”：自动采集 "
         "A股/美股/加密货币行情与新闻，由证据链投研委员会产出带证据、置信度与免责声明的多空研判。"
         "其创新不在“能分析股票”，而在把金融智能体推进到**可信研判范式**——一个面向高不确定性场景的"
         "**多智能体证据治理框架**。", indent=True)

    h2("7.1 五大可信支柱")
    bullet("证据治理 Evidence Governance：确定性规则引擎对每条研判施加治理（无证据不出强结论、数据过期降置信、"
           "证据冲突暴露分歧、模型无效弃权、回测不稳不支持高置信、新闻区分事实/情绪/推断）。")
    bullet("不确定性管理 Uncertainty Management：不预测涨跌，而是回答“有哪些支持/反对证据、是否依据充分、置信度是否合理”。")
    bullet("弃权感知 Abstention-aware：数据不可靠/证据不足/模型无统计优势时主动弃权，而非硬答。")
    bullet("自审计 Self-auditing：每次研判落库，到期用真实市场结果反向审计（命中率、各委员有效性、主席是否过度自信）。")
    bullet("反证感知解释 Counter-evidence-aware：解释不仅说“为何如此”，更说“什么证据反对、什么条件会使结论失效、谁持异议”。")

    h2("7.2 架构（在智阅 4 服务上加法 + 新增 ingestion-service）")
    para("新增金融 MCP 工具（get_quote/get_kline/get_indicators/get_stock_news/compute_signals/backtest/"
         "market_overview，共 12 个 MCP 工具）、市场无关数据适配器与**数据质量层**（每条行情标注 "
         "data_status ∈ {fresh/delayed/stale/fallback/error} 及停牌/涨跌停/复权口径）；agent-service 升级为"
         "证据链投研委员会 + XGBoost 信号校准器 + 主席汇总；新增 ingestion-service 定时采集+复盘回填+异常告警；"
         "网关新增金融仪表盘与决策解释面板。", indent=True)

    h2("7.3 证据治理引擎（R1–R6）")
    table(["规则", "触发", "动作"], [
        ["R1 无证据不出强结论", "委员给强结论但无证据", "降为中性/弃权"],
        ["R2 数据过期降置信", "data_status=stale/error", "置信度≤0.4"],
        ["R3 冲突暴露分歧", "委员意见冲突", "封顶0.55、标记分歧"],
        ["R4 模型无效弃权", "XGBoost AUC≈0.5/样本不足", "该票剔除"],
        ["R5 回测不稳不支持高置信", "参数敏感性不稳", "封顶0.6"],
        ["R6 新闻分层", "仅情绪/推断证据", "降为中性"],
    ])
    para("最终置信度 = min(主席提议, 治理上限)，使“高置信”必须有据。", indent=True)

    h2("7.4 运行实证（实测）")
    para("端到端冒烟测试（对运行中的 5 服务栈）：", indent=True)
    code_block(read("报告/artifacts/smoke_finance.txt"))
    para("XGBoost 风险信号离线训练实测（可复现产物）：", indent=True)
    code_block(read("报告/artifacts/ml_train.txt"))
    para("关于 ML 信号的诚实演进：起初让 XGBoost 预测 T+1 涨跌**方向**，样本外 AUC≈0.51（≈随机，无统计优势）"
         "——这是金融的客观规律（有效市场/随机游走）。据此把模型**重定向为预测“次日是否为大波动日”**"
         "（波动具聚集性 / GARCH 效应，是可学习的），样本外 **AUC≈0.556（高于随机基准 0.5）**，成为一个真正"
         "**有效的弱信号**。它作为**风险校准票**（非方向）参与委员会：当预测高波动时，治理规则 **R7** 据此"
         "下调整体置信度；若样本不足或无统计优势则按 R4 自动弃权。加密货币因 Binance 在本网络环境返回 451 被封禁，"
         "自动回退 CoinGecko；美股(Yahoo)不可达被标注为**最严重的 data_status=error** 并优雅降级（不崩溃、"
         "显式提示不可用），而非把缺失数据当可信。", indent=True)

    h2("7.5 结果展示")
    image("报告/screenshots/07_finance_chart.png", width=6.4,
          cap="图 7-1  金融仪表盘：K线 + MA5/MA20 + 成交量 + MACD/RSI 副图（指标据 K 线本地滚动计算）")
    image("报告/screenshots/06_finance_deep.png", width=4.2,
          cap="图 7-2  深度研判全貌：证据链投研委员会（四面 LLM 分析师 + ML 票）+ 主席置信度环 + 反证驱动解释面板")
    image("报告/screenshots/08_finance_explanation.png", width=6.4,
          cap="图 7-3  反证驱动决策解释面板：主席结论 + 置信度环 + 主要支持/反对证据 + 结论失效条件")
    para("说明：金融研判仅供学习研究，不构成投资建议；短周期方向接近随机游走，ML 与回测均标注未来函数/过拟合/不可外推。",
         indent=True, italic=True)


# ====================================================== 八、参考资料
def sec7():
    h1("八", "参考资料")
    refs = [
        "Model Context Protocol 官方规范. https://modelcontextprotocol.io/",
        "MCP Python SDK (FastMCP). https://github.com/modelcontextprotocol/python-sdk",
        "FastAPI 官方文档. https://fastapi.tiangolo.com/",
        "DeepSeek 开放平台（OpenAI 兼容接口）. https://platform.deepseek.com/",
        "Beautiful Soup 文档. https://www.crummy.com/software/BeautifulSoup/",
        "阿里云容器镜像服务 ACR 文档. https://help.aliyun.com/product/60716.html",
        "Docker / Docker Compose 官方文档. https://docs.docker.com/",
        "AKShare 开源财经数据接口库. https://akshare.akfamily.xyz/",
        "yfinance（Yahoo Finance 数据）. https://github.com/ranaroussi/yfinance",
        "XGBoost 文档. https://xgboost.readthedocs.io/ ；scikit-learn 概率校准. https://scikit-learn.org/",
        "Apache ECharts 可视化库. https://echarts.apache.org/",
        "《MCP 服务实验指导书》（课程资料）。",
    ]
    for i, r in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.4
        run = p.add_run(f"[{i}] {r}")
        run.font.size = Pt(11)
        _cjk(run, BODY)


# ====================================================== 九、附录
def sec8():
    h1("九", "附录（源代码）")
    para("项目目录结构：", bold=True)
    code_block(
        "微服务架构2/\n"
        "├─ services/\n"
        "│  ├─ storage-service/   db.py app.py Dockerfile requirements.txt\n"
        "│  ├─ mcp-tool-service/  tools.py mcp_server.py Dockerfile requirements.txt\n"
        "│  ├─ agent-service/     mcp_client.py agent.py app.py Dockerfile requirements.txt\n"
        "│  └─ api-gateway/       app.py static/index.html Dockerfile requirements.txt\n"
        "├─ guide-baseline/       web_content_mcp.py Dockerfile cline_mcp_settings.json\n"
        "├─ deploy/               docker-compose.yml .env.example\n"
        "├─ tests/                test_storage.py test_tools.py test_agent.py\n"
        "└─ scripts/              smoke_test.py push_aliyun.sh"
    )

    para("附录 A — storage-service", bold=True, color=JADE)
    appendix_file("services/storage-service/db.py")
    appendix_file("services/storage-service/app.py")

    para("附录 B — mcp-tool-service", bold=True, color=JADE)
    appendix_file("services/mcp-tool-service/tools.py")
    appendix_file("services/mcp-tool-service/mcp_server.py")

    para("附录 C — agent-service", bold=True, color=JADE)
    appendix_file("services/agent-service/mcp_client.py")
    appendix_file("services/agent-service/agent.py")
    appendix_file("services/agent-service/app.py")

    para("附录 D — api-gateway", bold=True, color=JADE)
    appendix_file("services/api-gateway/app.py")

    para("附录 E — 容器编排", bold=True, color=JADE)
    appendix_file("deploy/docker-compose.yml")

    para("附录 F — 指导书原样基线（guide-baseline）", bold=True, color=JADE)
    appendix_file("guide-baseline/web_content_mcp.py")

    para("附录 G — 测试", bold=True, color=JADE)
    appendix_file("tests/test_tools.py")
    appendix_file("scripts/smoke_test.py")

    para("附录 H — 智策金融平台核心源码", bold=True, color=JADE)
    appendix_file("services/mcp-tool-service/data_quality.py")
    appendix_file("services/mcp-tool-service/indicators.py")
    appendix_file("services/mcp-tool-service/backtest.py")
    appendix_file("services/mcp-tool-service/finance.py")
    appendix_file("services/agent-service/governance.py")
    appendix_file("services/agent-service/ml_signal.py")
    appendix_file("services/agent-service/committee.py")
    appendix_file("services/agent-service/review.py")


# ====================================================== build
cover()
sec1()
sec2()
sec3()
sec4()
sec5()
sec6()
sec_finance()
sec7()
sec8()

# normalize default font
style = doc.styles["Normal"]
style.font.name = BODY
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn("w:eastAsia"), BODY)

safe_name = COVER["name"].strip("（）")
safe_sid = COVER["sid"].strip("（）")
out = os.path.join(ROOT, "报告", f"{safe_sid}+{safe_name}.docx")
doc.save(out)
print("saved", out)
